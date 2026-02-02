import streamlit as st
import pandas as pd
from openai import OpenAI
import math
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import time
from datetime import datetime

# -----------------------------------------------------------------------------
# 1. APP CONFIGURATION & SESSION STATE
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Teach+ Pro", page_icon="🏫", layout="wide")

# Initialize Session State for Auth and Database Simulation
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_role' not in st.session_state:
    st.session_state['user_role'] = None
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'submitted_plans' not in st.session_state:
    st.session_state['submitted_plans'] = [] # Simulates a database

# -----------------------------------------------------------------------------
# 2. HELPER FUNCTIONS
# -----------------------------------------------------------------------------
@st.cache_data
def load_data():
    try:
        df = pd.read_csv("Teachshank_Master_Database_FINAL.tsv", sep='\t')
        return df
    except FileNotFoundError:
        return pd.DataFrame()

def calculate_pacing(df):
    if df.empty: return pd.DataFrame()
    groups = df.groupby('Chapter Name')['Learning Outcomes'].count().reset_index()
    groups.columns = ['Chapter', 'Topic_Count']
    total = groups['Topic_Count'].sum()
    if total > 0:
        groups['Allocated_Days'] = (groups['Topic_Count'] / total) * 180
        groups['Allocated_Days'] = groups['Allocated_Days'].apply(lambda x: max(1, math.ceil(x)))
    return groups

def create_word_docx(lesson_text, metadata):
    doc = Document()
    title = doc.add_heading('Teach+ Lesson Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    r1 = table.rows[0].cells
    r1[0].paragraphs[0].add_run(f"Teacher: {metadata['teacher']}").bold = True
    r1[1].paragraphs[0].add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}").bold = True
    
    r2 = table.rows[1].cells
    r2[0].paragraphs[0].add_run(f"Class: {metadata['grade']} | Subject: {metadata['subject']}\nUnit: {metadata['chapter']}")
    r2[1].paragraphs[0].add_run(f"Lesson: {metadata['lesson_num']} of {metadata['total_days']}\nStatus: {metadata['status']}")
    
    doc.add_paragraph("") 

    for line in lesson_text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('## '):
            h = doc.add_heading(line.replace('## ', ''), level=1)
            h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('### '):
            h = doc.add_heading(line.replace('### ', ''), level=2)
            h.runs[0].font.color.rgb = RGBColor(51, 102, 153)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line.replace('- ', '').replace('* ', ''), style='List Bullet')
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_strict_lesson(grade, subject, chapter, outcome, duration, lesson_num, total_days, api_key):
    progress_ratio = lesson_num / total_days
    if lesson_num == 1:
        phase = "PHASE 1: ENGAGE (Intro)"
        instr = "Day 1. Focus ONLY on sparking curiosity."
    elif lesson_num == total_days:
        phase = "PHASE 4: EVALUATE (Assessment)"
        instr = "Final Day. Summative assessment only."
    elif progress_ratio <= 0.5:
        phase = "PHASE 2: EXPLORE (Deep Dive)"
        instr = "Core instruction. Focus on 'How' and 'Why'."
    else:
        phase = "PHASE 3: ELABORATE (Application)"
        instr = "Application day. Practice and real-life scenarios."

    prompt = f"""
    You are the 'Teach+' Strict Pacing Engine.
    CONTEXT: {grade} - {subject} | Unit: {chapter} | Day {lesson_num} of {total_days}
    PHASE: {phase} | FOCUS: {outcome}
    INSTRUCTION: {instr}
    
    OUTPUT (Markdown):
    ## 1. Lesson Metadata
    - **Lesson Position:** {lesson_num}/{total_days} ({phase})
    - **Goal:** (Specific to this phase)

    ## 2. Micro-Script ({duration} Mins)
    (Provide strict 5E script based on phase)
    """
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o", messages=[{"role": "user", "content": prompt}], temperature=0.5
        )
        return response.choices[0].message.content, phase
    except Exception as e:
        return f"Error: {str(e)}", "Error"

# -----------------------------------------------------------------------------
# 3. LOGIN SCREEN
# -----------------------------------------------------------------------------
def login_screen():
    st.markdown("<h1 style='text-align: center;'>🔐 Teach+ Secure Login</h1>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Login")
            
            if submit:
                # HARDCODED CREDENTIALS FOR DEMO (In prod, use st.secrets or DB)
                if username == "teacher" and password == "teach123":
                    st.session_state['logged_in'] = True
                    st.session_state['user_role'] = "Teacher"
                    st.session_state['username'] = "Ms. Riya (Teacher)"
                    st.rerun()
                elif username == "admin" and password == "admin123":
                    st.session_state['logged_in'] = True
                    st.session_state['user_role'] = "Principal"
                    st.session_state['username'] = "Dr. Sharma (Principal)"
                    st.rerun()
                else:
                    st.error("Invalid credentials. Try 'teacher/teach123' or 'admin/admin123'")

# -----------------------------------------------------------------------------
# 4. DASHBOARD (LOGGED IN VIEW)
# -----------------------------------------------------------------------------
def main_dashboard():
    # Sidebar Profile
    st.sidebar.markdown(f"👤 **{st.session_state['username']}**")
    if st.sidebar.button("Logout"):
        st.session_state['logged_in'] = False
        st.rerun()
    st.sidebar.markdown("---")

    df = load_data()
    if df.empty:
        st.error("Database not found on GitHub.")
        return

    # --- PRINCIPAL VIEW: REVIEW DASHBOARD ---
    if st.session_state['user_role'] == "Principal":
        st.title("👨‍🏫 Principal's Desk")
        st.info("Review lesson plans submitted by teachers.")
        
        tab_review, tab_planner = st.tabs(["📂 Review Queue", "🛠️ Curriculum Planner"])
        
        with tab_review:
            if not st.session_state['submitted_plans']:
                st.write("✅ No pending submissions.")
            else:
                for i, plan in enumerate(st.session_state['submitted_plans']):
                    with st.expander(f"{plan['date']} | {plan['teacher']} | {plan['chapter']} (Day {plan['lesson_num']})"):
                        st.markdown(plan['content'])
                        c1, c2 = st.columns(2)
                        if c1.button(f"👍 Approve Plan {i+1}", key=f"app_{i}"):
                            st.success(f"Plan approved!")
                            # In real app, update DB here
                        if c2.button(f"👎 Request Changes {i+1}", key=f"rej_{i}"):
                            st.warning("Feedback sent to teacher.")
    
    # --- TEACHER VIEW (DEFAULT) ---
    else:
        st.title("📝 Teacher's Planner")
        tab_planner = st.container()

    # --- SHARED PLANNER LOGIC ---
    # Both Teacher and Principal can access the generator
    
    with tab_planner if st.session_state['user_role'] == "Principal" else st.container():
        st.header("1️⃣ Select Class & Subject")
        grade = st.selectbox("Grade", sorted(df['Grade'].astype(str).unique()))
        subject = st.selectbox("Subject", sorted(df[df['Grade'].astype(str)==grade]['Subject'].unique()))
        
        subject_data = df[(df['Grade'].astype(str)==grade) & (df['Subject']==subject)]
        pacing_df = calculate_pacing(subject_data)
        
        # Two Tabs: Calendar & Designer
        t1, t2 = st.tabs(["📅 Annual Pacing", "✍️ Lesson Designer"])
        
        with t1:
            if not pacing_df.empty:
                st.dataframe(pacing_df, use_container_width=True, 
                             column_config={"Allocated_Days": st.column_config.ProgressColumn("Budget (Days)", max_value=30)})
        
        with t2:
            c1, c2 = st.columns([1, 2])
            with c1:
                chapters = pacing_df['Chapter'].tolist()
                if not chapters: st.stop()
                sel_chapter = st.selectbox("Unit", chapters)
                
                total_days = int(pacing_df[pacing_df['Chapter'] == sel_chapter]['Allocated_Days'].values[0])
                st.caption(f"📅 Budget: **{total_days} Days**")
                
                lesson_num = st.number_input(f"Lesson (1-{total_days})", 1, total_days, 1)
                
                outcomes = subject_data[subject_data['Chapter Name'] == sel_chapter]['Learning Outcomes'].tolist()
                sel_outcome = st.selectbox("Topic", outcomes)
                duration = st.slider("Mins", 30, 90, 45)
                
                if st.button("🚀 Generate Plan", type="primary"):
                    if "OPENAI_API_KEY" not in st.secrets:
                        st.error("API Key Missing")
                    else:
                        with st.spinner("Generating..."):
                            script, phase = generate_strict_lesson(
                                grade, subject, sel_chapter, sel_outcome, duration, 
                                lesson_num, total_days, st.secrets["OPENAI_API_KEY"]
                            )
                            # Store in session state for the "Submit" phase
                            st.session_state['current_plan'] = {
                                "script": script,
                                "phase": phase,
                                "meta": {
                                    "grade": grade, "subject": subject, "chapter": sel_chapter,
                                    "topic": sel_outcome, "duration": duration, 
                                    "lesson_num": lesson_num, "total_days": total_days, 
                                    "phase": phase, "teacher": st.session_state['username'],
                                    "status": "Draft"
                                }
                            }

            with c2:
                if 'current_plan' in st.session_state:
                    plan = st.session_state['current_plan']
                    st.subheader(f"📄 Draft: {plan['meta']['chapter']} (Day {plan['meta']['lesson_num']})")
                    st.markdown(plan['script'])
                    
                    # ACTION BUTTONS
                    ac1, ac2 = st.columns(2)
                    
                    # 1. DOWNLOAD
                    docx = create_word_docx(plan['script'], plan['meta'])
                    ac1.download_button("📥 Download Word Doc", docx, "LessonPlan.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
                    # 2. SUBMIT TO PRINCIPAL
                    if st.session_state['user_role'] == "Teacher":
                        if ac2.button("📤 Submit to Principal"):
                            submission_entry = {
                                "teacher": st.session_state['username'],
                                "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                                "chapter": plan['meta']['chapter'],
                                "lesson_num": plan['meta']['lesson_num'],
                                "content": plan['script']
                            }
                            st.session_state['submitted_plans'].append(submission_entry)
                            st.success("✅ Submitted successfully to Principal's Dashboard!")
                            time.sleep(2)
                            st.rerun()

# -----------------------------------------------------------------------------
# 5. APP ENTRY POINT
# -----------------------------------------------------------------------------
if st.session_state['logged_in']:
    main_dashboard()
else:
    login_screen()
